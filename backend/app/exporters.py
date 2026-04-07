# backend/app/exporters.py
import io
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_docx(title: str, content: str, style: str) -> bytes:
    """Создание DOCX файла в памяти"""
    doc = Document()
    
    # Заголовок документа
    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Метаданные
    meta = doc.add_paragraph()
    meta.add_run(f"Стиль: {style}").italic = True
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Отступ
    
    # Разбиваем контент на абзацы
    paragraphs = content.split('\n\n')
    
    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
            
        # Проверяем, является ли абзац заголовком (короткий, без точки в конце, часто с :)
        if len(para) < 100 and not para.endswith('.') and (':' in para or para.isupper() or para.startswith('Введение') or para.startswith('Заключение') or para.startswith('Раздел')):
            doc.add_heading(para, level=2)
        else:
            # Обычный абзац
            p = doc.add_paragraph(para)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Сохраняем в байты
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()
